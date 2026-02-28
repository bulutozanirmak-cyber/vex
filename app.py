import streamlit as st
import os
import re
import zipfile
import textwrap
import html
from docx import Document
from docx.shared import Pt
import fitz
from fpdf import FPDF
import io

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="VEX | Profesyonel Hukuk Dönüştürücü", page_icon="⚖️")

st.title("⚖️ VEX: Master Architect")
st.markdown("---")

# --- FONT TANIMLAMALARI ---
FONT_REG = "Roboto-Regular.ttf"
FONT_BOLD = "Roboto-Bold.ttf"
FONT_ITAL = "Roboto-Italic.ttf"

class VexWebApp:
    @staticmethod
    def decode_text(data):
        for encoding in ['utf-8', 'windows-1254', 'iso-8859-9', 'cp1254']:
            try: return data.decode(encoding)
            except: continue
        return data.decode('utf-8', errors='ignore')

    @staticmethod
    def extract_udf_data(raw_xml):
        """Veri çekme ve boşluk koruma motoru [cite: 32-36]."""
        text = html.unescape(raw_xml)
        cdata_text = ""
        if "<![CDATA[" in text:
            cdata_text = text.split("<![CDATA[")[1].split("]]>")[0]
        
        if not cdata_text.strip() or "il_Ilce" in cdata_text[:100]:
            deep_text = re.sub(r'<[^>]+>', '\n', text) 
            deep_text = re.sub(r'\n\s*\n', '\n\n', deep_text)
            return deep_text.strip()
        return cdata_text

    @staticmethod
    def save_pdf_pro(content):
        """AttributeError hatası giderilmiş, fpdf2 uyumlu PDF motoru."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_margins(20, 15, 20)
        
        # Font Yükleme Mantığı
        try:
            if os.path.exists(FONT_REG):
                pdf.add_font('Roboto', '', FONT_REG)
                pdf.add_font('Roboto', 'B', FONT_BOLD) if os.path.exists(FONT_BOLD) else None
                pdf.add_font('Roboto', 'I', FONT_ITAL) if os.path.exists(FONT_ITAL) else None
                pdf.set_font('Roboto', '', 11)
            else:
                pdf.set_font('Helvetica', '', 10)
        except:
            pdf.set_font('Helvetica', '', 10)

        wrapper = textwrap.TextWrapper(width=80, break_long_words=True, replace_whitespace=False)
        for line in content.split('\n'):
            if not line.strip():
                pdf.ln(7) # Satır boşluklarını koru
            else:
                for s_line in wrapper.wrap(line):
                    pdf.cell(0, 6, text=s_line, ln=1, align='L')
        
        # --- KRİTİK DÜZELTME ---
        # fpdf2'de output() doğrudan bytes döndürür, .encode() gerekmez.
        return bytes(pdf.output())

    @staticmethod
    def save_udf_pro(content):
        lines = content.split('\n')
        p_xml, offset = "", 0
        for l in lines:
            l_len = len(l) + 1
            p_xml += f'<paragraph><content startOffset="{offset}" length="{l_len}" /></paragraph>\n'
            offset += l_len
        
        xml = f"""<?xml version="1.0" encoding="UTF-8" ?>
<template format_id="1.8"><content><![CDATA[{content}]]></content>
<properties><pageFormat mediaSizeName="1" leftMargin="42.525" rightMargin="42.525" topMargin="42.525" bottomMargin="42.525" paperOrientation="1" headerFOffset="20.0" footerFOffset="20.0" /></properties>
<elements resolver="hvl-default">{p_xml}</elements>
<styles><style name="default" family="Times New Roman" size="12" /><style name="hvl-default" family="Times New Roman" size="12" /></styles></template>"""
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as z:
            z.writestr("content.xml", xml)
        return zip_buffer.getvalue()

    @staticmethod
    def save_docx_pro(content):
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        return doc_io.getvalue()

# --- ARAYÜZ MANTIĞI ---
uploaded_files = st.file_uploader("Dosyaları yükleyin", accept_multiple_files=True)

if uploaded_files:
    vex_files = []
    for uploaded_file in uploaded_files:
        bytes_data = uploaded_file.read()
        ext = os.path.splitext(uploaded_file.name)[1].lower()
        
        try:
            if ext == ".pdf":
                doc = fitz.open(stream=bytes_data, filetype="pdf")
                text = "\n".join([p.get_text("text", sort=True) for p in doc])
            elif ext == ".udf":
                with zipfile.ZipFile(io.BytesIO(bytes_data)) as udf:
                    with udf.open('content.xml') as f:
                        raw = VexWebApp.decode_text(f.read())
                        text = VexWebApp.extract_udf_data(raw)
            elif ext == ".docx":
                doc = Document(io.BytesIO(bytes_data))
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                text = VexWebApp.decode_text(bytes_data)
            
            vex_files.append({"name": uploaded_file.name, "content": text})
        except Exception as e:
            st.error(f"Hata: {e}")

    if vex_files:
        st.success(f"{len(vex_files)} dosya yüklendi.")
        target_format = st.selectbox("Format Seçin", ["PDF", "UDF", "DOCX", "TXT"])

        if st.button("Dönüştür ve İndir"):
            for f in vex_files:
                if target_format == "PDF":
                    data = VexWebApp.save_pdf_pro(f['content'])
                    st.download_button(f"📥 {f['name']} (PDF)", data, f"VEX_{f['name'][:2]}.pdf")
                elif target_format == "UDF":
                    data = VexWebApp.save_udf_pro(f['content'])
                    st.download_button(f"📥 {f['name']} (UDF)", data, f"VEX_{f['name'][:2]}.udf")
                elif target_format == "DOCX":
                    data = VexWebApp.save_docx_pro(f['content'])
                    st.download_button(f"📥 {f['name']} (DOCX)", data, f"VEX_{f['name'][:2]}.docx")
